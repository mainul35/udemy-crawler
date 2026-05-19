from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document


def build(input_json: Path, output_docx: Path) -> tuple[int, int]:
    data = json.loads(input_json.read_text(encoding="utf-8"))
    quizzes = data["quizzes"]

    for quiz in quizzes:
        if not quiz.get("title"):
            raise ValueError(f"Quiz index={quiz.get('index')!r} is missing a title")
        if not quiz.get("questions"):
            raise ValueError(f"Quiz index={quiz.get('index')!r} has no questions")

    doc = Document()

    for quiz in quizzes:
        for q_num, question in enumerate(quiz["questions"], start=1):
            doc.add_heading(f"Quiz {quiz['index']} — {quiz['title']}", level=2)
            doc.add_paragraph(f"{q_num}. {question['stem']}")
            for option in question["options"]:
                doc.add_paragraph(f"   {option['letter']}. {option['text']}")
            doc.add_page_break()

    doc.add_heading("Answer Key", level=1)
    for quiz in quizzes:
        doc.add_heading(f"Quiz {quiz['index']} — {quiz['title']}", level=2)
        for q_num, question in enumerate(quiz["questions"], start=1):
            correct = ", ".join(o["letter"] for o in question["options"] if o["is_correct"])
            doc.add_paragraph(f"Q{q_num}: {correct}")

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx))

    total_questions = sum(len(q["questions"]) for q in quizzes)
    return len(quizzes), total_questions


def main() -> int:
    input_json = Path(sys.argv[1])
    output_docx = Path(sys.argv[2])
    n_quizzes, n_questions = build(input_json, output_docx)
    size = output_docx.stat().st_size
    print(f"OK path={output_docx} bytes={size} quizzes={n_quizzes} questions={n_questions}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
