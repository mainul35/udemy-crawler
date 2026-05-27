"""Build a Word document with each question's correct option highlighted.

Same layout as ``_build_doc.py`` (one question per page, answer key at the
end) but the correct option line is rendered with a yellow background and
bold text so it stands out at a glance.

Usage:
    python scripts/_build_doc_highlighted.py <quizzes.json> <output.docx>
"""
from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX


def _add_option(doc, option: dict) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"   {option['letter']}. {option['text']}")
    if option["is_correct"]:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        run.bold = True


def build(input_json: Path, output_docx: Path) -> tuple[int, int]:
    data = json.loads(input_json.read_text(encoding="utf-8"))
    quizzes = data["quizzes"]

    for quiz in quizzes:
        if not quiz.get("title"):
            raise ValueError(f"Quiz index={quiz.get('index')!r} is missing a title")
        if not quiz.get("questions"):
            raise ValueError(f"Quiz index={quiz.get('index')!r} has no questions")

    doc = Document()

    for quiz in quizzes:
        for q_num, question in enumerate(quiz["questions"], start=1):
            doc.add_heading(f"Quiz {quiz['index']} — {quiz['title']}", level=2)
            doc.add_paragraph(f"{q_num}. {question['stem']}")
            for option in question["options"]:
                _add_option(doc, option)
            doc.add_page_break()

    doc.add_heading("Answer Key", level=1)
    for quiz in quizzes:
        doc.add_heading(f"Quiz {quiz['index']} — {quiz['title']}", level=2)
        for q_num, question in enumerate(quiz["questions"], start=1):
            correct = ", ".join(o["letter"] for o in question["options"] if o["is_correct"])
            doc.add_paragraph(f"Q{q_num}: {correct}")

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx))

    total_questions = sum(len(q["questions"]) for q in quizzes)
    return len(quizzes), total_questions


def main() -> int:
    input_json = Path(sys.argv[1])
    output_docx = Path(sys.argv[2])
    n_quizzes, n_questions = build(input_json, output_docx)
    size = output_docx.stat().st_size
    print(f"OK path={output_docx} bytes={size} quizzes={n_quizzes} questions={n_questions}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
